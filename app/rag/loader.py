import os
from typing import List
from langchain_core.documents import Document
try:
    from langchain_text_splitters import RecursiveCharacterTextSplitter
except ImportError:
    from langchain.text_splitter import RecursiveCharacterTextSplitter

from app.utils.logger import get_logger

logger = get_logger("RAG.Loader")

def load_pdf(file_path: str) -> str:
    """Extracts text from a PDF file using pypdf or pdfplumber."""
    text = ""
    try:
        import pypdf
        reader = pypdf.PdfReader(file_path)
        for page in reader.pages:
            t = page.extract_text()
            if t:
                text += t + "\n"
    except Exception as e:
        logger.warning(f"pypdf failed on {file_path}: {e}. Trying pdfplumber...")
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    t = page.extract_text()
                    if t:
                        text += t + "\n"
        except Exception as ex:
            logger.error(f"Failed to read PDF {file_path}: {ex}")
    return text

def load_docx(file_path: str) -> str:
    """Extracts text from a Word DOCX file."""
    text = ""
    try:
        import docx
        doc = docx.Document(file_path)
        for paragraph in doc.paragraphs:
            if paragraph.text:
                text += paragraph.text + "\n"
    except Exception as e:
        logger.error(f"Failed to read DOCX {file_path}: {e}")
    return text

def load_txt(file_path: str) -> str:
    """Extracts text from a raw TXT file."""
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception as e:
        logger.error(f"Failed to read TXT {file_path}: {e}")
        return ""

def process_file(file_path: str, chunk_size: int = 800, chunk_overlap: int = 150) -> List[Document]:
    """
    Ingests PDF, DOCX, or TXT file and splits it into Document chunks with metadata.
    """
    filename = os.path.basename(file_path)
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        raw_text = load_pdf(file_path)
    elif ext in [".docx", ".doc"]:
        raw_text = load_docx(file_path)
    elif ext in [".txt", ".md", ".csv"]:
        raw_text = load_txt(file_path)
    else:
        logger.warning(f"Unsupported file format: {ext}")
        return []

    if not raw_text.strip():
        logger.warning(f"No text extracted from file: {filename}")
        return []

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", ". ", " ", ""]
    )
    
    chunks = splitter.split_text(raw_text)
    documents = [
        Document(
            page_content=chunk,
            metadata={
                "source": filename,
                "file_path": file_path,
                "chunk_index": idx,
                "total_chunks": len(chunks)
            }
        )
        for idx, chunk in enumerate(chunks)
    ]
    logger.info(f"Ingested '{filename}': Created {len(documents)} document chunks.")
    return documents
